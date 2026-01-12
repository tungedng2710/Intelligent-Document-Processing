from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import httpx
import os
import uuid
import json
import logging
from pathlib import Path
from typing import Optional
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from bs4 import BeautifulSoup
import shutil

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Document Analysis Webapp")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
ANALYSIS_SERVICE_URL = os.getenv("ANALYSIS_SERVICE_URL", "http://0.0.0.0:9890/parse")
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/")
async def read_root():
    """Serve the main HTML page"""
    return FileResponse("static/index.html")


@app.post("/api/analyze")
async def analyze_document(file: UploadFile = File(...)):
    """
    Upload a file and send it to the analysis service
    """
    # Generate unique filename
    file_id = str(uuid.uuid4())
    file_extension = Path(file.filename).suffix
    temp_file_path = UPLOAD_DIR / f"{file_id}{file_extension}"
    
    logger.info(f"Analyzing file: {file.filename}")
    logger.info(f"Sending to service: {ANALYSIS_SERVICE_URL}")
    
    try:
        # Save uploaded file
        with open(temp_file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Send file to analysis service
        async with httpx.AsyncClient(timeout=120.0) as client:
            with open(temp_file_path, "rb") as f:
                files = {"file": (file.filename, f, file.content_type)}
                response = await client.post(ANALYSIS_SERVICE_URL, files=files)
        
        if response.status_code != 200:
            raise HTTPException(
                status_code=response.status_code,
                detail=f"Analysis service returned error: {response.text}"
            )
        
        result = response.json()
        
        # Clean up temporary file
        if temp_file_path.exists():
            os.remove(temp_file_path)
        
        # Extract markdown and output_image from response
        markdown_content = result.get("markdown", "")
        output_image = result.get("output_image", None)
        
        response_data = {
            "success": True,
            "markdown": markdown_content,
            "file_id": file_id,
            "original_filename": file.filename
        }
        
        # Include output_image if present (for detection mode)
        if output_image:
            response_data["output_image"] = output_image
        
        return JSONResponse(response_data)
    
    except httpx.TimeoutException:
        if temp_file_path.exists():
            os.remove(temp_file_path)
        raise HTTPException(status_code=504, detail="Analysis service timeout")
    
    except Exception as e:
        if temp_file_path.exists():
            os.remove(temp_file_path)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/convert-to-docx")
async def convert_to_docx(markdown_text: str = File(...), filename: Optional[str] = None):
    """
    Convert markdown text to DOCX format
    """
    try:
        # Generate unique filename for DOCX
        file_id = str(uuid.uuid4())
        base_filename = filename.replace(Path(filename).suffix, "") if filename else "document"
        docx_filename = f"{base_filename}_{file_id}.docx"
        docx_path = UPLOAD_DIR / docx_filename
        
        # Create DOCX document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Convert markdown to HTML first for better parsing
        html_content = markdown.markdown(markdown_text, extensions=['extra', 'codehilite', 'tables'])
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Parse HTML and add to document
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'pre', 'table']):
            if element.name.startswith('h'):
                # Headers
                level = int(element.name[1])
                para = doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                # Paragraphs
                para = doc.add_paragraph(element.get_text())
            elif element.name in ['ul', 'ol']:
                # Lists
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
            elif element.name == 'pre':
                # Code blocks
                para = doc.add_paragraph(element.get_text())
                para.style = 'No Spacing'
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            elif element.name == 'table':
                # Tables
                rows = element.find_all('tr')
                if rows:
                    cols = len(rows[0].find_all(['th', 'td']))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            table.rows[i].cells[j].text = cell.get_text()
        
        # If simple parsing fails, fallback to plain text with basic formatting
        if len(doc.paragraphs) == 0:
            lines = markdown_text.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
        
        # Save document
        doc.save(docx_path)
        
        return FileResponse(
            path=docx_path,
            filename=docx_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            background=lambda: os.remove(docx_path) if docx_path.exists() else None
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error converting to DOCX: {str(e)}")


@app.post("/api/download-markdown")
async def download_markdown(markdown_text: str = File(...), filename: Optional[str] = None):
    """
    Download markdown text as .md file
    """
    try:
        # Generate unique filename for markdown
        file_id = str(uuid.uuid4())
        base_filename = filename.replace(Path(filename).suffix, "") if filename else "document"
        md_filename = f"{base_filename}_{file_id}.md"
        md_path = UPLOAD_DIR / md_filename
        
        # Save markdown content
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
        
        return FileResponse(
            path=md_path,
            filename=md_filename,
            media_type="text/markdown",
            background=lambda: os.remove(md_path) if md_path.exists() else None
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error downloading markdown: {str(e)}")


@app.get("/api/service-urls")
async def get_service_urls():
    """Get available service URLs from config file"""
    try:
        service_urls_file = Path("service_urls.json")
        if service_urls_file.exists():
            with open(service_urls_file, 'r') as f:
                service_urls = json.load(f)
            return {"service_urls": service_urls, "current": ANALYSIS_SERVICE_URL}
        else:
            return {"service_urls": {"0.0.0.0:7875/parse": "Default Service"}, "current": ANALYSIS_SERVICE_URL}
    except Exception as e:
        return {"service_urls": {"0.0.0.0:7875/parse": "Default Service"}, "current": ANALYSIS_SERVICE_URL}


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get(ANALYSIS_SERVICE_URL.replace('/parse', '/health'))
            service_status = "online" if response.status_code == 200 else "offline"
    except:
        service_status = "offline"
    
    return {
        "status": "online",
        "analysis_service": service_status,
        "analysis_service_url": ANALYSIS_SERVICE_URL
    }


@app.post("/api/update-service-url")
async def update_service_url(request: Request):
    """Update the analysis service URL"""
    global ANALYSIS_SERVICE_URL
    try:
        body = await request.json()
        new_url = body.get("new_url")
        if not new_url:
            raise HTTPException(status_code=400, detail="new_url is required")
        
        # Ensure URL has http:// prefix
        if not new_url.startswith('http'):
            new_url = f"http://{new_url}"
        
        ANALYSIS_SERVICE_URL = new_url
        logger.info(f"Service URL updated to: {ANALYSIS_SERVICE_URL}")
        return {"success": True, "new_url": ANALYSIS_SERVICE_URL}
    except Exception as e:
        logger.error(f"Error updating service URL: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=7869)
